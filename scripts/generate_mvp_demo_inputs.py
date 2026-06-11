from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


OUTPUT_DIR = Path("demo_inputs")


EMAIL_CASES = {
    "email_tariff_reclassification.txt": """From: broker.ops@pacific-customs.example
Sent: Thursday, June 11, 2026 08:40 AM
To: import-compliance@greenchem.example
Subject: URGENT HTS Review - China-Origin Electronics Components SHP-CHN-4481

Shipment SHP-CHN-4481 is inbound from Shanghai (CNSHA) to Los Angeles (USLAX) on Pacific Star V-118.
SKU affected: SKU-EL-4892, industrial electronics controller boards.
Declared value: 1,180,000 USD.

The customs broker believes the current HTS code 8537.10 may be incomplete and may require review under a Section 301 China tariff line.
Estimated additional duty exposure is 142,000 USD if reclassified.

Please open a compliance review ticket and notify finance before entry summary filing.
""",
    "email_missing_coo_customs_hold.txt": """From: release.team@border-broker.example
Sent: Thursday, June 11, 2026 10:15 AM
To: logistics@greenchem.example
Subject: CBP Hold - Missing Country of Origin Certificate for SHP-VNM-2099

CBP has placed shipment SHP-VNM-2099 on document hold at Port of Los Angeles.
Origin port: Ho Chi Minh City (VNSGN).
Destination port: Los Angeles (USLAX).
SKU affected: SKU-TEXT-8842, synthetic industrial poly-weave rolls.

The commercial invoice says Vietnam origin, but the packing list references China-origin raw material.
The supplier has not provided a country of origin certificate or substantial transformation support.
Delivery appointment is within 48 hours. Demurrage may begin tomorrow.

Please alert logistics and request corrected COO documentation from the supplier immediately.
""",
    "email_tsca_chemical_import.txt": """From: supplier.docs@chem-sourcing.example
Sent: Thursday, June 11, 2026 11:05 AM
To: compliance@greenchem.example
Subject: Missing TSCA Certification - Solvent Additive Shipment SHP-CHEM-7720

Shipment SHP-CHEM-7720 contains solvent additive SKU-CHEM-1102 for import into the United States.
Origin port: Busan (KRPUS).
Destination port: Long Beach (USLGB).

The supplier provided an SDS but omitted the CAS number for one component and did not provide a TSCA import certification statement.
The customs broker asked whether the substance is subject to EPA TSCA requirements or qualifies for negative certification.

Please hold release until compliance confirms CAS numbers, SDS completeness, and TSCA import status.
""",
    "email_hazmat_lithium_battery.txt": """From: airfreightdesk@carrier.example
Sent: Thursday, June 11, 2026 13:20 PM
To: inbound-ops@greenchem.example
Subject: Dangerous Goods Documentation Gap - Lithium Battery Modules

Air shipment AIR-LAX-3307 includes lithium battery modules, SKU-BATT-5519.
Origin airport: Shenzhen (SZX).
Destination airport: Los Angeles (LAX).

The booking references UN3480 lithium ion batteries, but the dangerous goods declaration is missing packing instruction details.
The package label photo shows Class 9 hazard marking, but the shipping papers do not include emergency response information.

Carrier will not tender cargo until PHMSA/DOT hazardous materials documentation is verified.
""",
    "email_uflpa_supplier_risk.txt": """From: sourcing.audit@greenchem.example
Sent: Thursday, June 11, 2026 15:10 PM
To: legal-compliance@greenchem.example
Subject: UFLPA Traceability Concern - Solar Controller Supplier

Supplier Xinjiang-linked component concern identified for shipment SHP-SOLAR-6610.
Product: solar inverter controller assemblies, SKU-SOLAR-2201.
Raw material disclosure mentions polysilicon subcomponents from an upstream supplier that has not provided chain-of-custody documents.

The supplier declined to provide full traceability records and asked whether a purchase order summary would be sufficient.
Because the goods are inbound to the United States, please review UFLPA forced labor due diligence requirements before release.
""",
}


def write_email_cases() -> None:
    for filename, content in EMAIL_CASES.items():
        (OUTPUT_DIR / filename).write_text(content, encoding="utf-8")


def write_csv_case() -> None:
    rows = [
        {
            "shipment_id": "SHP-CHN-4481",
            "sku": "SKU-EL-4892",
            "origin_country": "China",
            "destination_port": "USLAX",
            "hts_code": "8537.10",
            "declared_value_usd": "1180000",
            "document_status": "HTS review required",
            "risk_notes": "Broker flagged possible Section 301 exposure",
        },
        {
            "shipment_id": "SHP-VNM-2099",
            "sku": "SKU-TEXT-8842",
            "origin_country": "Vietnam",
            "destination_port": "USLAX",
            "hts_code": "5903.90",
            "declared_value_usd": "420000",
            "document_status": "COO missing",
            "risk_notes": "Packing list references China-origin raw material",
        },
        {
            "shipment_id": "SHP-CHEM-7720",
            "sku": "SKU-CHEM-1102",
            "origin_country": "South Korea",
            "destination_port": "USLGB",
            "hts_code": "3814.00",
            "declared_value_usd": "260000",
            "document_status": "TSCA certification missing",
            "risk_notes": "CAS number missing for solvent component",
        },
    ]
    with (OUTPUT_DIR / "shipment_batch_risk_cases.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_json_case() -> None:
    payload = {
        "source": "simulated_port_congestion_api",
        "scraped_at": "2026-06-11T16:00:00Z",
        "port": "USLAX",
        "signal_type": "port_congestion",
        "severity": "HIGH",
        "affected_shipments": ["SHP-VNM-2099", "SHP-CHN-4481"],
        "details": {
            "average_dwell_days": 5.8,
            "demurrage_risk": True,
            "customs_documentation_backlog": True,
            "notes": "Terminal reports higher dwell time for containers pending CBP document review.",
        },
        "recommended_review": [
            "Prioritize shipments with missing COO documents",
            "Notify logistics if delivery deadline is within 72 hours",
            "Estimate demurrage and detention exposure",
        ],
    }
    (OUTPUT_DIR / "port_congestion_signal.json").write_text(
        json.dumps(payload, indent=2),
        encoding="utf-8",
    )


def write_pdf_case() -> None:
    path = OUTPUT_DIR / "customs_hold_notice_missing_coo.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("CBP DOCUMENT HOLD NOTICE", styles["Title"]),
        Spacer(1, 12),
    ]

    lines = [
        "Shipment ID: SHP-VNM-2099",
        "Carrier: Ocean Bridge Express",
        "Origin Port: Ho Chi Minh City (VNSGN)",
        "Destination Port: Los Angeles (USLAX)",
        "SKU Affected: SKU-TEXT-8842",
        "Issue: Country of origin certificate missing.",
        "Invoice origin states Vietnam, while packing list references China-origin raw materials.",
        "Delivery deadline: within 48 hours.",
        "Operational Risk: CBP hold may cause demurrage, storage fees, and missed delivery appointment.",
        "Requested Action: Obtain COO certificate and substantial transformation support from supplier.",
    ]

    for line in lines:
        story.append(Paragraph(line, styles["BodyText"]))
        story.append(Spacer(1, 6))

    doc.build(story)


def write_excel_case() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Import Risk Matrix"
    headers = [
        "shipment_id",
        "sku",
        "origin_country",
        "destination_port",
        "risk_type",
        "document_gap",
        "estimated_exposure_usd",
        "delivery_deadline",
        "recommended_action",
    ]
    rows = [
        ["SHP-CHN-4481", "SKU-EL-4892", "China", "USLAX", "Tariff", "HTS review", 142000, "2026-06-18", "JIRA compliance review"],
        ["SHP-CHEM-7720", "SKU-CHEM-1102", "South Korea", "USLGB", "TSCA", "CAS/TSCA certification missing", 35000, "2026-06-20", "Request SDS/CAS/TSCA statement"],
        ["AIR-LAX-3307", "SKU-BATT-5519", "China", "LAX", "Hazmat", "DG declaration incomplete", 58000, "2026-06-13", "Verify PHMSA/DOT shipping papers"],
    ]
    sheet.append(headers)
    for row in rows:
        sheet.append(row)

    header_fill = PatternFill("solid", fgColor="1F4E78")
    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill

    for column_cells in sheet.columns:
        max_length = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells)
        sheet.column_dimensions[column_cells[0].column_letter].width = min(max(max_length + 2, 12), 38)

    workbook.save(OUTPUT_DIR / "shipment_import_risk_matrix.xlsx")


def write_word_case() -> None:
    document = Document()
    document.add_heading("Supplier Origin and Forced Labor Attestation", level=1)
    document.add_paragraph("Supplier: North Ridge Solar Components Ltd.")
    document.add_paragraph("Related shipment: SHP-SOLAR-6610")
    document.add_paragraph("SKU affected: SKU-SOLAR-2201")
    document.add_heading("Supplier Statement", level=2)
    document.add_paragraph(
        "The supplier states that final assembly occurs outside China, but upstream "
        "polysilicon component traceability is incomplete. The supplier has not "
        "provided full chain-of-custody documents for all raw materials."
    )
    document.add_heading("Compliance Concern", level=2)
    document.add_paragraph(
        "Because the shipment is inbound to the United States and the product category "
        "may include solar or polysilicon components, legal/compliance should review "
        "potential UFLPA forced labor due diligence requirements before release."
    )
    document.add_heading("Requested Action", level=2)
    document.add_paragraph("Request supplier traceability records and notify legal/compliance.")
    document.save(OUTPUT_DIR / "supplier_origin_uflpa_attestation.docx")


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    write_email_cases()
    write_csv_case()
    write_json_case()
    write_pdf_case()
    write_excel_case()
    write_word_case()
    print(f"Generated MVP demo inputs in {OUTPUT_DIR.resolve()}")


if __name__ == "__main__":
    main()
